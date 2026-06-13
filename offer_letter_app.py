import streamlit as st
import zipfile
import io
import os
import subprocess
import tempfile
import shutil
from datetime import date, timedelta
import random
import string

try:
    from docx2pdf import convert as _docx2pdf_convert
    _HAS_DOCX2PDF = True
except ImportError:
    _HAS_DOCX2PDF = False

# ── Page config ──────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Harion Research – Offer Letter Generator",
    page_icon="📄",
    layout="centered",
)

# ── Branding / CSS ────────────────────────────────────────────────────────────
st.markdown("""
<style>
    /* Force dark theme */
    html, body, [data-testid="stAppViewContainer"], [data-testid="stApp"] {
        background-color: #0e1117 !important;
        color: #fafafa !important;
    }
    [data-testid="stHeader"] { background-color: #0e1117 !important; }
    [data-testid="stSidebar"] { background-color: #161b22 !important; }

    .stApp { background: #0e1117; }
    .main-card {
        background: #1c2333;
        border-radius: 14px;
        padding: 2.5rem 2.5rem 2rem;
        box-shadow: 0 2px 18px rgba(0,0,0,0.4);
        margin-top: 1.5rem;
        color: #fafafa;
    }
    .brand-header {
        display:flex; align-items:center; gap:14px;
        border-bottom: 2px solid #4a90d9; padding-bottom:1rem; margin-bottom:1.8rem;
    }
    .brand-title { font-size:1.55rem; font-weight:700; color:#4a90d9; margin:0; }
    .brand-sub   { font-size:.85rem; color:#8ab4d4; margin:0; }
    .section-label { font-size:.78rem; font-weight:600; color:#8ab4d4;
                     text-transform:uppercase; letter-spacing:.08em; margin-bottom:.3rem; }
    div[data-testid="stButton"] button {
        background:#4a90d9; color:white; border:none;
        border-radius:8px; padding:.65rem 2rem; font-size:1rem; font-weight:600;
        width:100%; margin-top:.5rem; cursor:pointer;
    }
    div[data-testid="stDownloadButton"] button {
        background:#0e6e3e; color:white; border:none;
        border-radius:8px; padding:.65rem 2rem; font-size:1rem; font-weight:600;
        width:100%; margin-top:.5rem;
    }
    .preview-box {
        background:#0e1117; border-radius:10px; padding:1.5rem 2rem;
        border-left: 4px solid #4a90d9; margin-top:1.2rem;
        font-size:.92rem; line-height:1.7; color:#fafafa;
    }
    .success-badge {
        background:#0e2e1e; color:#3dd68c; border-radius:6px;
        padding:.45rem 1rem; font-weight:600; font-size:.9rem;
        display:inline-block; margin-bottom:1rem;
    }

    /* Inputs */
    [data-testid="stTextInput"] input {
        background:#0e1117 !important; color:#fafafa !important;
        border: 1px solid #2d3a4a !important; border-radius:8px !important;
    }
    [data-testid="stDateInput"] input {
        background:#0e1117 !important; color:#fafafa !important;
        border: 1px solid #2d3a4a !important;
    }

    /* Tabs */
    [data-testid="stTabs"] button {
        color: #8ab4d4 !important;
    }
    [data-testid="stTabs"] button[aria-selected="true"] {
        color: #4a90d9 !important;
        border-bottom: 2px solid #4a90d9 !important;
    }

    /* Info / warning boxes */
    [data-testid="stAlert"] {
        background:#1c2333 !important; color:#fafafa !important;
    }
</style>
""", unsafe_allow_html=True)

EQUITY_TEMPLATE_PATH          = os.path.join(os.path.dirname(__file__), "offer_letter_temp.docx")
MARKETING_TEMPLATE_PATH       = os.path.join(os.path.dirname(__file__), "offer_letter_marketing_temp.docx")
EQUITY_CERT_TEMPLATE_PATH     = os.path.join(os.path.dirname(__file__), "Cert.docx")
MARKETING_CERT_TEMPLATE_PATH  = os.path.join(os.path.dirname(__file__), "Cert_marketing.docx")


# ── Helpers ───────────────────────────────────────────────────────────────────

def fill_docx_template(template_path: str, name: str, start_date: str, ref_no: str) -> bytes:
    """Replace {} placeholders in the docx XML, inject Ref No, return new docx bytes."""
    with open(template_path, "rb") as f:
        docx_bytes = f.read()

    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as zin:
        files = {n: zin.read(n) for n in zin.namelist()}

    xml = files["word/document.xml"].decode("utf-8")
    first  = xml.index("{}")
    xml    = xml[:first] + name + xml[first + 2:]
    second = xml.index("{}")
    xml    = xml[:second] + start_date + xml[second + 2:]
    files["word/document.xml"] = xml.encode("utf-8")

    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for n, data in files.items():
            zout.writestr(n, data)

    # Inject the Ref No paragraph after the title (paragraph index 2)
    return inject_ref_no(out.getvalue(), ref_no, insert_after_idx=2)


def _find_libreoffice() -> str | None:
    for binary in ["soffice", "libreoffice"]:
        path = shutil.which(binary)
        if path:
            return path
    common_paths = [
        "/usr/bin/soffice",
        "/usr/bin/libreoffice",
        "/usr/lib/libreoffice/program/soffice",
        "/opt/libreoffice/program/soffice",
        "/snap/bin/libreoffice",
    ]
    for p in common_paths:
        if os.path.isfile(p):
            return p
    return None


def docx_to_pdf(docx_bytes: bytes) -> tuple[bytes | None, str]:
    """Convert docx bytes to PDF bytes.
    Returns (pdf_bytes, error_msg). pdf_bytes is None on failure.
    Priority: docx2pdf (MS Word COM via subprocess) → LibreOffice → None.
    """
    import sys as _sys
    with tempfile.TemporaryDirectory() as tmp:
        docx_path = os.path.join(tmp, "document.docx")
        pdf_path  = os.path.join(tmp, "document.pdf")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        # ── 1. Try Word COM via DispatchEx (always a fresh Word process) ──────
        try:
            helper = os.path.join(tmp, "_conv_helper.py")
            with open(helper, "w") as _hf:
                _hf.write(
                    "import sys, os\n"
                    "from pathlib import Path\n"
                    "import win32com.client\n"
                    "docx = str(Path(sys.argv[1]).resolve())\n"
                    "pdf  = str(Path(sys.argv[2]).resolve())\n"
                    "word = win32com.client.DispatchEx('Word.Application')\n"
                    "word.Visible = False\n"
                    "try:\n"
                    "    doc = word.Documents.Open(docx)\n"
                    "    doc.SaveAs(pdf, FileFormat=17)\n"
                    "    doc.Close(0)\n"
                    "finally:\n"
                    "    word.Quit()\n"
                )
            result = subprocess.run(
                [_sys.executable, helper, docx_path, pdf_path],
                capture_output=True, text=True, timeout=180,
            )
            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                with open(pdf_path, "rb") as f:
                    return f.read(), ""
            err = (result.stderr or result.stdout or "no output").strip()
            print("[Word COM] failed:", err)
        except subprocess.TimeoutExpired:
            print("[docx2pdf] timed out")
            err = "docx2pdf timed out after 180 s"
        except Exception as e:
            err = str(e)
            print("[docx2pdf] Error:", err)

        # ── 2. Fall back to LibreOffice ───────────────────────────────────
        lo_binary = _find_libreoffice()
        if lo_binary:
            lo_profile = os.path.join(tmp, "lo_profile")
            os.makedirs(lo_profile, exist_ok=True)
            env = os.environ.copy()
            env["HOME"]   = lo_profile
            env["TMPDIR"] = tmp
            profile_url = "file://" + lo_profile
            cmd = [
                lo_binary,
                "-env:UserInstallation=" + profile_url,
                "--headless", "--norestore", "--nofirststartwizard",
                "--convert-to", "pdf",
                "--outdir", tmp,
                docx_path,
            ]
            try:
                lo_result = subprocess.run(cmd, capture_output=True, text=True, timeout=180, env=env)
                lo_pdf = docx_path.replace(".docx", ".pdf")
                if os.path.exists(lo_pdf):
                    with open(lo_pdf, "rb") as f:
                        return f.read(), ""
                err += " | LO: " + (lo_result.stderr or "no output").strip()
            except Exception as e:
                err += " | LO error: " + str(e)

    return None, err


def ordinal(n: int) -> str:
    suffix = {1: "st", 2: "nd", 3: "rd"}.get(
        n % 10 if n % 100 not in (11, 12, 13) else 0, "th"
    )
    return f"{n}{suffix}"


def gen_ref_no() -> str:
    """Generate a unique reference number like HAR/A3X9/7BKQ."""
    chars = string.ascii_uppercase + string.digits
    part1 = ''.join(random.choices(chars, k=4))
    part2 = ''.join(random.choices(chars, k=4))
    return f"HAR/{part1}/{part2}"


def inject_ref_no(docx_bytes: bytes, ref_no: str, insert_after_idx: int = 2) -> bytes:
    """Insert a Ref No paragraph into the document using python-docx.

    Copies the run formatting from the paragraph at insert_after_idx so the
    font, size, and colour match the surrounding text.
    insert_after_idx=2 works for both offer-letter and certificate templates
    (the blank line after the main heading).

    For certificate templates: also removes 2 of the 3 empty paragraphs before
    the signing block to prevent the content overflowing onto a second page.
    """
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy as _copy

    doc = Document(io.BytesIO(docx_bytes))
    paras = doc.paragraphs

    # Determine a reference paragraph to copy formatting from
    ref_para = paras[min(insert_after_idx + 1, len(paras) - 1)]

    # Build a new paragraph element
    new_para_elem = OxmlElement('w:p')

    # Copy paragraph properties (alignment, spacing) from the reference para
    if ref_para._p.pPr is not None:
        new_para_elem.append(_copy.deepcopy(ref_para._p.pPr))

    # Build a run with the ref-no text
    r_elem = OxmlElement('w:r')

    # Copy run properties from the first run of the reference paragraph if any
    if ref_para.runs:
        src_rpr = ref_para.runs[0]._r.find(qn('w:rPr'))
        if src_rpr is not None:
            r_elem.append(_copy.deepcopy(src_rpr))

    t_elem = OxmlElement('w:t')
    t_elem.text = f"Ref No.: {ref_no}"
    t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r_elem.append(t_elem)
    new_para_elem.append(r_elem)

    # Insert after the paragraph at insert_after_idx
    anchor = paras[insert_after_idx]._p
    anchor.addnext(new_para_elem)

    # For certificate templates: trim 2 of the 3 empty paragraphs before the
    # signing block so everything fits on one page.
    # Detect cert by looking for "For Harion Research" paragraph.
    paras = doc.paragraphs  # refresh after insert
    signing_idx = next(
        (i for i, p in enumerate(paras) if p.text.strip().startswith("For Harion Research")),
        None,
    )
    if signing_idx is not None:
        # Collect consecutive empty paragraphs immediately before the signing block
        empty_before = []
        j = signing_idx - 1
        while j >= 0 and paras[j].text.strip() == "":
            empty_before.append(paras[j]._p)
            j -= 1
        # Remove all but 1 empty paragraph (keep 1 for spacing)
        for p_elem in empty_before[1:]:
            p_elem.getparent().remove(p_elem)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def fill_cert_docx_template(
    template_path: str,
    name: str,
    from_date: str,
    to_date: str,
    issue_date: str,
    ref_no: str,
    pronoun: str = "him",
) -> bytes:
    """Replace the {} placeholders in the certificate template.

    Confirmed placeholder order (from XML inspection):
      1st  {} → issue date      (after "Date:")
      2nd  {} → intern name     ("certify that {} has successfully")
      3rd  {} → from date       ("from {} to")
      4th  {} → to date         ("to {}")
      5th  {} → intern name     ("put in by {}")
      6th  {} → pronoun         (first occurrence, e.g. "him" / "her")
      7th  {} → pronoun         (second occurrence, e.g. "him" / "her")
    """
    with open(template_path, "rb") as f:
        docx_bytes = f.read()

    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as zin:
        files = {n: zin.read(n) for n in zin.namelist()}

    xml = files["word/document.xml"].decode("utf-8")
    for replacement in [issue_date, name, from_date, to_date, name, pronoun, pronoun]:
        idx = xml.index("{}")
        xml = xml[:idx] + replacement + xml[idx + 2:]
    files["word/document.xml"] = xml.encode("utf-8")

    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for n, data in files.items():
            zout.writestr(n, data)

    # Inject Ref No paragraph after "TO WHOMSOEVER IT MAY CONCERN" (index 2)
    return inject_ref_no(out.getvalue(), ref_no, insert_after_idx=2)


def render_certificate_form(role_title: str, template_path: str, key_prefix: str):
    """Reusable certificate generation form."""
    st.markdown(
        f"Fill in the details below to generate a personalised **{role_title}** "
        "completion certificate as a PDF."
    )

    # Row 1 - name
    st.markdown('<p class="section-label">Intern Full Name</p>', unsafe_allow_html=True)
    cert_name = st.text_input(
        "Intern Full Name",
        placeholder="e.g. Ananya Sharma",
        label_visibility="collapsed",
        key=f"{key_prefix}_cert_name",
    )

    # Row 2 - gender selector
    st.markdown('<p class="section-label">Intern Gender (pronoun)</p>', unsafe_allow_html=True)
    gender_choice = st.radio(
        "Gender",
        options=["Male (him/his)", "Female (her/her)"],
        horizontal=True,
        label_visibility="collapsed",
        key=f"{key_prefix}_gender",
    )
    pronoun = "him" if gender_choice.startswith("Male") else "her"

    # Row 3 - internship period
    col_from, col_to = st.columns(2)
    with col_from:
        st.markdown('<p class="section-label">Internship From</p>', unsafe_allow_html=True)
        cert_from = st.date_input(
            "From",
            value=date.today().replace(day=1),
            label_visibility="collapsed",
            format="DD/MM/YYYY",
            key=f"{key_prefix}_cert_from",
        )
    with col_to:
        st.markdown('<p class="section-label">Internship To</p>', unsafe_allow_html=True)
        cert_to = st.date_input(
            "To",
            value=date.today(),
            label_visibility="collapsed",
            format="DD/MM/YYYY",
            key=f"{key_prefix}_cert_to",
        )

    # Row 4 - date of issuance
    st.markdown('<p class="section-label">Date of Issuance</p>', unsafe_allow_html=True)
    cert_issue = st.date_input(
        "Date of Issuance",
        value=date.today(),
        label_visibility="collapsed",
        format="DD/MM/YYYY",
        key=f"{key_prefix}_cert_issue",
    )

    fmt_from  = f"{ordinal(cert_from.day)} {cert_from.strftime('%B %Y')}"
    fmt_to    = f"{ordinal(cert_to.day)} {cert_to.strftime('%B %Y')}"
    fmt_issue = f"{ordinal(cert_issue.day)} {cert_issue.strftime('%B %Y')}"

    # Live preview
    if cert_name.strip():
        st.markdown(
            f'<div class="preview-box"><strong>Preview snippet:</strong><br><br>'
            f'<span style="color:#8ab4d4;font-size:.82rem;">Ref No.: HAR/XXXX/XXXX &nbsp;(generated on download)</span><br><br>'
            f'Date of Issuance: <strong>{fmt_issue}</strong><br><br>'
            f'This is to certify that <strong>{cert_name}</strong> has successfully completed '
            f'the internship as a <strong>{role_title}</strong> at <strong>Harion Research</strong> '
            f'from <strong>{fmt_from}</strong> to <strong>{fmt_to}</strong>.<br><br>'
            f'We wish <strong>{pronoun}</strong> all the best in <strong>{pronoun}</strong> future endeavours.</div>',
            unsafe_allow_html=True,
        )
    else:
        st.info("\U0001f446 Enter the intern\u2019s name above to see a live preview.")

    st.markdown("<br>", unsafe_allow_html=True)

    # Session-state keys - persist generated bytes across Streamlit reruns
    ss_pdf  = f"{key_prefix}_pdf_bytes"
    ss_docx = f"{key_prefix}_docx_bytes"
    ss_name = f"{key_prefix}_safe_name"

    if st.button("\U0001f393 Generate Certificate PDF", key=f"{key_prefix}_cert_btn"):
        if not cert_name.strip():
            st.error("Please enter the intern's name before generating.")
        elif cert_to < cert_from:
            st.error("'To' date must be on or after 'From' date.")
        else:
            ref_no = gen_ref_no()
            with st.spinner("Filling certificate template and converting to PDF\u2026"):
                docx_bytes = fill_cert_docx_template(
                    template_path,
                    cert_name.strip(),
                    fmt_from,
                    fmt_to,
                    fmt_issue,
                    ref_no,
                    pronoun,
                )
                pdf_bytes, pdf_err = docx_to_pdf(docx_bytes)

            safe_name = cert_name.strip().replace(" ", "_")
            # Store in session state so the download button survives the rerun
            st.session_state[ss_pdf]  = pdf_bytes
            st.session_state[ss_docx] = docx_bytes
            st.session_state[ss_name] = safe_name
            st.session_state[ss_pdf + "_err"] = pdf_err

    # Render download button OUTSIDE the generate block so it survives reruns
    if st.session_state.get(ss_pdf):
        st.markdown(
            '<p class="success-badge">\u2705 Certificate ready for download!</p>',
            unsafe_allow_html=True,
        )
        st.download_button(
            label="\u2b07\ufe0f  Download Certificate PDF",
            data=st.session_state[ss_pdf],
            file_name=f"Certificate_{st.session_state[ss_name]}.pdf",
            mime="application/pdf",
            key=f"{key_prefix}_cert_dl_pdf",
        )
    elif st.session_state.get(ss_docx):
        err_msg = st.session_state.get(ss_pdf + "_err", "")
        st.warning(f"PDF conversion failed — downloading as Word document instead.\n\n`{err_msg}`" if err_msg else "PDF conversion failed — downloading as Word document instead.")
        st.download_button(
            label="\u2b07\ufe0f  Download Certificate DOCX",
            data=st.session_state[ss_docx],
            file_name=f"Certificate_{st.session_state[ss_name]}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"{key_prefix}_cert_dl_docx",
        )


def render_offer_letter_form(role_title: str, template_path: str, key_prefix: str):
    """Reusable form for any offer letter type."""
    st.markdown(f"Fill in the candidate details below to generate a personalised **{role_title}** offer letter as a PDF.")

    col1, col2 = st.columns([3, 2])

    with col1:
        st.markdown('<p class="section-label">Candidate Full Name</p>', unsafe_allow_html=True)
        candidate_name = st.text_input(
            "Candidate Full Name",
            placeholder="e.g. Ananya Sharma",
            label_visibility="collapsed",
            key=f"{key_prefix}_name"
        )

    with col2:
        st.markdown('<p class="section-label">Internship Start Date</p>', unsafe_allow_html=True)
        start_date = st.date_input(
            "Start Date",
            value=date.today() + timedelta(days=7),
            min_value=date.today(),
            label_visibility="collapsed",
            format="DD/MM/YYYY",
            key=f"{key_prefix}_date"
        )

    formatted_date = f"{ordinal(start_date.day)} {start_date.strftime('%B %Y')}"

    if candidate_name.strip():
        st.markdown(f"""
        <div class="preview-box"><strong>Preview snippet:</strong><br><br>
<span style="color:#8ab4d4;font-size:.82rem;">Ref No.: HAR/XXXX/XXXX &nbsp;(generated on download)</span><br><br>
Dear <strong>{candidate_name}</strong>,<br><br>
On behalf of <strong>Harion Research</strong>, I am pleased to offer you the position of
<strong>{role_title}</strong> for a duration of <strong>2 months</strong>,
starting from <strong>{formatted_date}</strong>.
        </div>""", unsafe_allow_html=True)
    else:
        st.info("👆 Enter the candidate's name above to see a live preview.")

    st.markdown("<br>", unsafe_allow_html=True)

    if st.button("⚡ Generate Offer Letter PDF", key=f"{key_prefix}_btn"):
        if not candidate_name.strip():
            st.error("Please enter the candidate's name before generating.")
        else:
            ref_no = gen_ref_no()
            with st.spinner("Filling template and converting to PDF\u2026"):
                docx_bytes = fill_docx_template(template_path, candidate_name.strip(), formatted_date, ref_no)
                pdf_bytes, pdf_err = docx_to_pdf(docx_bytes)

            safe_name = candidate_name.strip().replace(" ", "_")

            if pdf_bytes:
                st.markdown('<p class="success-badge">\u2705 Offer letter ready for download!</p>',
                            unsafe_allow_html=True)
                st.download_button(
                    label="\u2b07\ufe0f  Download PDF",
                    data=pdf_bytes,
                    file_name=f"Offer_Letter_{safe_name}.pdf",
                    mime="application/pdf",
                    key=f"{key_prefix}_dl_pdf"
                )
            else:
                st.warning(
                    f"PDF conversion failed \u2014 downloading as Word document instead.\n\n`{pdf_err}`"
                    if pdf_err else
                    "PDF conversion failed \u2014 downloading as Word document instead."
                )
                st.download_button(
                    label="\u2b07\ufe0f  Download DOCX",
                    data=docx_bytes,
                    file_name=f"Offer_Letter_{safe_name}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"{key_prefix}_dl_docx"
                )


# ── Email sender ─────────────────────────────────────────────────────────────

SMTP_HOST   = "smtp.zoho.in"
SMTP_PORT   = 465
FROM_EMAIL  = "hr@harionresearch.co.in"


def _get_zoho_password() -> str:
    """Return Zoho password from secrets.toml or environment, else empty string."""
    try:
        pw = st.secrets.get("ZOHO_PASSWORD", "")
        if pw and pw != "your-zoho-app-password-here":
            return pw
    except Exception:
        pass
    return os.environ.get("ZOHO_PASSWORD", "")


def render_email_form():
    """Simple email composer that sends via Zoho SMTP."""
    import smtplib
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    from email.mime.base import MIMEBase
    from email import encoders as _enc

    st.markdown(
        "Send an email directly from "
        "**hr@harionresearch.co.in** via Zoho Mail."
    )

    # ── Password ──────────────────────────────────────────────────────────────
    stored_pw = _get_zoho_password()
    if not stored_pw:
        st.info(
            "🔑 No app password found in `secrets.toml`. "
            "Enter it below — it won't be stored."
        )
    zoho_pw = st.text_input(
        "Zoho App Password",
        value=stored_pw,
        type="password",
        placeholder="Paste your Zoho app password",
        help="Generate one in Zoho Mail → Settings → Security → App Passwords",
        key="email_zoho_pw",
    )

    st.markdown("---")

    # ── To ────────────────────────────────────────────────────────────────────
    st.markdown('<p class="section-label">To — Recipient Email Address</p>', unsafe_allow_html=True)
    to_email = st.text_input(
        "To",
        placeholder="e.g. candidate@gmail.com",
        label_visibility="collapsed",
        help="The email address of the person you are writing to.",
        key="email_to",
    )

    # ── Subject ───────────────────────────────────────────────────────────────
    st.markdown('<p class="section-label">Subject</p>', unsafe_allow_html=True)
    subject = st.text_input(
        "Subject",
        placeholder="e.g. Offer Letter – Harion Research Internship",
        label_visibility="collapsed",
        key="email_subject",
    )

    # ── Body ──────────────────────────────────────────────────────────────────
    st.markdown('<p class="section-label">Message Body</p>', unsafe_allow_html=True)
    body = st.text_area(
        "Body",
        placeholder="Type your email message here…",
        height=220,
        label_visibility="collapsed",
        key="email_body",
    )

    # ── Attachment ────────────────────────────────────────────────────────────
    st.markdown('<p class="section-label">Attachment (optional)</p>', unsafe_allow_html=True)
    attachments = st.file_uploader(
        "Attach files",
        label_visibility="collapsed",
        accept_multiple_files=True,
        help="Attach one or more PDF, DOCX, or any files to send with the email.",
        key="email_attachment",
    )
    if attachments:
        for f in attachments:
            st.caption(f"📎 {f.name}  ({round(len(f.getvalue()) / 1024, 1)} KB)")

    st.markdown("<br>", unsafe_allow_html=True)

    # ── Send ──────────────────────────────────────────────────────────────────
    if st.button("📤 Send Email", key="send_email_btn"):
        errors = []
        if not to_email.strip():
            errors.append("Recipient email address is required.")
        elif "@" not in to_email:
            errors.append("Recipient email address looks invalid.")
        if not subject.strip():
            errors.append("Subject is required.")
        if not body.strip():
            errors.append("Message body cannot be empty.")
        if not zoho_pw.strip():
            errors.append("Zoho app password is required to send.")

        if errors:
            for e in errors:
                st.error(e)
        else:
            with st.spinner("Connecting to Zoho and sending…"):
                try:
                    msg = MIMEMultipart()
                    msg["From"]    = FROM_EMAIL
                    msg["To"]      = to_email.strip()
                    msg["Subject"] = subject.strip()
                    msg.attach(MIMEText(body, "plain", "utf-8"))

                    for attachment in (attachments or []):
                        part = MIMEBase("application", "octet-stream")
                        part.set_payload(attachment.getvalue())
                        _enc.encode_base64(part)
                        part.add_header(
                            "Content-Disposition",
                            f'attachment; filename="{attachment.name}"',
                        )
                        msg.attach(part)

                    with smtplib.SMTP_SSL(SMTP_HOST, SMTP_PORT) as server:
                        server.login(FROM_EMAIL, zoho_pw.strip())
                        server.sendmail(FROM_EMAIL, to_email.strip(), msg.as_string())

                    st.success(f"✅ Email sent to **{to_email.strip()}** successfully!")
                    st.balloons()
                except smtplib.SMTPAuthenticationError:
                    st.error(
                        "❌ Authentication failed. Check your Zoho App Password "
                        "(not your login password — generate one in Zoho Settings → Security)."
                    )
                except smtplib.SMTPException as exc:
                    st.error(f"❌ SMTP error: {exc}")
                except Exception as exc:
                    st.error(f"❌ Unexpected error: {exc}")


# ── UI ────────────────────────────────────────────────────────────────────────

st.markdown("""
<div class="main-card">
  <div class="brand-header">
    <div>
      <p class="brand-title">📄 Harion Research</p>
      <p class="brand-sub">Internship Offer Letter Generator · HR Portal</p>
    </div>
  </div>
""", unsafe_allow_html=True)

tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "📊 Offer – Equity Research",
    "📣 Offer – Marketing",
    "🎓 Certificate – Equity Research",
    "🎓 Certificate – Marketing",
    "📧 Send Email",
])

with tab1:
    render_offer_letter_form(
        role_title="Equity Research Analyst Intern",
        template_path=EQUITY_TEMPLATE_PATH,
        key_prefix="equity",
    )

with tab2:
    render_offer_letter_form(
        role_title="Marketing Intern",
        template_path=MARKETING_TEMPLATE_PATH,
        key_prefix="marketing",
    )

with tab3:
    render_certificate_form(
        role_title="Equity Research Analyst Intern",
        template_path=EQUITY_CERT_TEMPLATE_PATH,
        key_prefix="cert_equity",
    )

with tab4:
    render_certificate_form(
        role_title="Marketing Intern",
        template_path=MARKETING_CERT_TEMPLATE_PATH,
        key_prefix="cert_marketing",
    )

with tab5:
    render_email_form()

st.markdown("</div>", unsafe_allow_html=True)

st.markdown("---")
st.markdown(
    "<center style='color:#aaa;font-size:.78rem;'>© 2025 Harion Research · HR Internal Tool</center>",
    unsafe_allow_html=True
)
